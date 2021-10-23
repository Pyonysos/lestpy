"""
TO DO:
-Create a tool "outlier detection"
|-- Add Mahalanobis distance
|-- Add z-score

-Translate the interaction in english

-implement decorator for time measurement

-Use black package to improve code readibility and respect the pythonic style of coding

"""


"""
DEPENDENCIES
"""

from pandas.core.frame import DataFrame
from sklearn import model_selection
from sklearn.preprocessing import RobustScaler, MinMaxScaler, StandardScaler
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score
from sklearn.model_selection import LeaveOneOut

from mpl_toolkits import mplot3d

from scipy.stats import dirichlet

import statsmodels.api as sm
from statsmodels.stats import outliers_influence

import inspect

#for print_in_file()
from docx import Document
from docx.shared import Inches

import matplotlib.pyplot as plt
import plotly.figure_factory as ff
import plotly.express as px
import pandas as pd
import numpy as np

from SALib.sample import saltelli
from SALib.analyze import sobol

import time

def Plot_surface(a,b,c, **kwargs):
    
    if hasattr(kwargs, "cmap"):
        cmap=kwargs['cmap']
    else:
        cmap='Viridis'
  
    fig = plt.figure(figsize=(14,9))    
    ax = plt.axes(projection='3d')

    surf = ax.plot_trisurf(a.ravel(), b.ravel(), c, cmap=cmap, antialiased=True, edgecolor='none')
            
    fig.colorbar(surf, ax =ax, shrink=0.5, aspect=5)
    ax.set_xlabel(f'{a.name}')
    ax.set_ylabel(f'{b.name}')
    ax.set_zlabel(f'{c.name}')
    ax.set_title(f'Interaction: {c.name}')
    fig.show()
    return

#interaction_dict will store all the data concerning the use of interactions
interaction_dict={}

"""
Definition of the interactions
"""
class Interaction:
    """
    Interaction class

    interactions are mathematical functions that aim to describe the logical interactions that can occur between two parameters on an output value.
    the LBM_Regression will calculate all the interactions of 2 features possible with the input features.
    Then it will build model using the fewest number of features or their interactions that best explains the response.
    As the interactions describes real physical effects, the user have the control to exclude interactions that are not relevant in their case study.
    """
    def __init__(self, x, y, max_x=None, min_x=None, max_y=None, min_y=None):
        self.x = x
        self.y = y
        if max_x == None:
            self.max_x = np.max(self.x)
        else:
            self.max_x = max_x
        if max_y == None:
            self.max_y = np.max(self.y)
        else:
            self.max_y = max_y
        if min_x == None:
            self.min_x = np.min(self.x)
        else:
            self.min_x = min_x
        if min_y == None:
            self.min_y = np.min(self.y)
        else:
            self.min_y = min_y

        if not hasattr(self.x, 'name') or not hasattr(self.y, 'name'):
            self.x = pd.DataFrame(self.x, name='x')
            self.y = pd.DataFrame(self.y, name='y')
     
    def compute(self):
        """
        compute the interaction
        return a dataframe named according to the interaction, with its values
        """
        return pd.DataFrame(self.calc(), columns=[self.name])
        
    def display_interaction(self, x=None, y=None):
        """
        display_interaction is a method to help visualize how the interaction is modeled by ifs function. if x and y are not given, it creates vectors of hundred numbers between -1 and 1 and calculates the values of the interaction.
        x : pandas dataframe, values of feature x
        y : pandas dataframe, values of feature y
        
        plot a surface of the interaction on x and y
        
        return None
        """
        if x is None:
            x= np.linspace(-1, 1, 100)
        if y is None:
            y= np.linspace(-1, 1, 100)
            
        x, y = np.meshgrid(x, y)
        x=pd.DataFrame(x.ravel(), columns=["x"])
        y=pd.DataFrame(y.ravel(), columns=["y"])
        z = self.calc()
        Plot_surface(x,y,z)
      

class X_fort_Quand_Y_faible_Et_Inversement(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} fort quand {self.y.name} faible et inversement'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array(np.multiply(-self.x, self.y)).reshape(-1,1)
        return func
        
class X_fort_Ou_Y_fort(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} fort ou {self.y.name} fort'
        self.interaction = self.__class__.__name__       
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array(-(self.max_x-self.x)*(self.max_y-self.y)).reshape(-1,1)
        return func
    
class X_fort_Ou_Y_faible(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} fort ou {self.y.name} faible'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
    
    def calc(self):
        func = np.array(-(self.max_x-self.x)*(np.abs(self.min_y)+self.y)).reshape(-1,1)
        return func

class X_et_Y_forts(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} et {self.y.name} forts'
        self.interaction = self.__class__.__name__ 
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array((self.x+np.abs(self.min_x))*(self.y+np.abs(self.min_y))).reshape(-1,1)
        return func
                                        
class X_fort_et_Y_faible(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} fort et {self.y.name} faible'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array((self.x+np.abs(self.min_x))*(self.max_y-self.y)).reshape(-1,1)
        return func

class X_fort_si_Y_fort(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} fort si {self.y.name} fort'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array(self.x*(np.abs(self.min_y)+self.y)).reshape(-1,1)
        return func
                                          
class X_fort_si_Y_faible(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} fort si {self.y.name} faible'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array(self.x*(np.abs(self.max_y)-self.y)).reshape(-1,1)
        return func
         
class X_fort_si_Y_moyen(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} fort si {self.y.name} moyen'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
    
    def calc(self):
        func = np.array(self.x / np.sqrt((self.max_y+np.abs(self.min_y))/500+np.square(self.y))).reshape(-1,1)
        return func
        
class X_moyen_si_Y_fort(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} moyen si {self.y.name} fort'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array((np.abs(self.min_y)+self.y)/np.sqrt((self.max_x+np.abs(self.min_x))/200+np.square(self.x))).reshape(-1,1)
        return func

class X_moyen_si_Y_faible(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} moyen si {self.y.name} faible'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array((self.max_y-self.y)/np.sqrt((self.max_x+np.abs(self.min_x))/200+np.square(self.x))).reshape(-1,1)
        return func
    
class Ni_X_ni_Y_extremes(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'Ni {self.x.name} ni {self.y.name} extremes'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array(-np.square(self.x)-np.square(self.y)).reshape(-1,1)
        return func
    
class X_Y_moyens(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} et {self.y.name} moyens'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array((np.square(self.max_x)-np.square(self.x))*(np.square(self.max_y)-np.square(self.y))).reshape(-1,1)
        return func
    
class X_comme_Y(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'{self.x.name} comme {self.y.name}'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array(-np.square(self.x-self.y)).reshape(-1,1)
        return func
    
class Somme_X_et_Y_forte(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'Somme {self.x.name} et {self.y.name} forte'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array(self.x+self.y).reshape(-1,1)
        return func
    
class Difference_X_et_Y_forte(Interaction):
    def __init__(self, x, y, max_x, min_x, max_y, min_y):
        super().__init__(x, y, max_x, min_x, max_y, min_y)
        self.name = f'Difference {self.x.name} et {self.y.name} forte'
        self.interaction = self.__class__.__name__
        interaction_dict[self.name] = {'x' : self.x, 'y' : self.y, 'interaction' : self.interaction}
        
    def calc(self):
        func = np.array(self.x-self.y).reshape(-1,1)
        return func

"""
**************************************************************************************************************************
REGRESSION ALGORHITHM
**************************************************************************************************************************
"""

class LBM_Regression:
    """
    Lesty Buat-Menard Regression.
        LBM_Regression calculate interactions of two variables, selects the most relevant ones and fits a linear model with coefficients w = (w1, ..., wp)
        to minimize the residual sum of squares between the observed targets in
        the dataset, and the targets predicted by the linear approximation.
    """
    def __init__(self):
        self.with_optimization = False
        self.with_interactions = False
        self.with_fit = False
        self.with_transform = False
        self.with_variable_instant = False
        return
    
    def bibliography(self):
        print("""
                Bibliography on which is based the project:
                1. Lesty, Michel, et P Buat-Ménard. "La synthèse géométrique des corrélations multidimensionnelles". Les Cahiers de l’Analyse des données VII, no 3 (1982): 355‑70.
                2. Lesty, Michel. "Une nouvelle approche dans le choix des régresseurs de la régression multiple en présence d’interactions et de colinéarités". revue Modulad 22 (1999): 41‑77.
                3. Derringer, George and Suich, Ronald. "Simultaneous Optimization of Several Response Variables". Journal of Quality Technology 12 (1980): 214-219. 
                """)    
    
    def __compute_interaction(self, X, allow_autointeraction, interaction_list):
        new_X = X.reset_index(drop=True)
        n = 1 if allow_autointeraction==True else 0
        for i in range(X.shape[1]):
            for j in range(i+1-n, X.shape[1]):
                for interaction in interaction_list:
                    if i != j:
                        #changer en numpy
                        new_X = pd.concat((new_X, eval(interaction)(X.iloc[:,i], X.iloc[:,j], X.iloc[:,i].max(axis=0),X.iloc[:,i].min(axis=0), X.iloc[:,j].max(axis=0), X.iloc[:,j].min(axis=0) ).compute()), axis =1)
                    else:
                        if not interaction in [Difference_X_et_Y_forte, X_comme_Y]:
                            new_X = pd.concat((new_X, eval(interaction)(X.iloc[:,i], X.iloc[:,j], X.iloc[:,i].max(axis=0),X.iloc[:,i].min(axis=0), X.iloc[:,j].max(axis=0), X.iloc[:,j].min(axis=0)).compute()), axis =1)
        
        self.with_interactions = True
                            
        return new_X
    
    def __rescale_data(self, X, y, scaler: str):
        """
        __rescale_data:
        scale the features with the specified method
        
        params:
        X: dataframe, dataframe of features
        y: dataframe, dataframe of response. useless to specify. for convenience only
        scaler: str, transformer to use to scale the data before the calculation of interactions
        
        returns: object
        """
        #normalisation des données
        if scaler =='robust':
            self.transformer = RobustScaler()
        elif scaler =='minmax':
            self.transformer = MinMaxScaler()
        elif scaler =='standard':
            self.transformer = StandardScaler()
        else:
          raise ValueError(f"{scaler} is not implemented. please use robust, standard or minmax")
        
        return self.transformer.fit_transform(X, y)
    
    def __variable_instant_transform(self, X, coef, denomin,L=None):
        """
        transformation of __variable_instant 
        """
        #print(Coef.shape, denomin.shape, X.shape, self.Shape)
        if L is None:
            L =[n for n in range(0, len(denomin))]
        
        """
        Coef = self.Coef[:,L] 
        """
        Denomin = pd.DataFrame(denomin).iloc[L].transpose()
        Denominator = np.array(Denomin).reshape(1, Denomin.shape[1])
        
        #print(Coef.shape, X.shape, self.Shape)
        return np.divide(np.subtract(np.array(X)*self.Shape, coef[:,L]), Denominator)
    
    def __variable_instant(self, X):
        """
        transformation of the matrix into a modified unit called "variable/instant"
        params:
        returns:
        """
        eye = np.eye(X.shape[0], X.shape[0])

        X_sqr = np.sum(np.multiply(X, X), axis=0)
        X_sum = np.square(np.sum(X, axis=0))
        self.denomin = np.sqrt(X.shape[0]*X_sqr-X_sum)*np.sqrt(X.shape[0]-1)

        self.Coef = np.sum(eye, axis=0).reshape(-1,1).dot(np.sum(np.array(X), axis=0).reshape(1,-1))[0]
        self.Coef = np.array(self.Coef).reshape(1, self.Coef.shape[0])
        
        self.with_variable_instant = True
        self.Shape = X.shape[0]
        
        #variable = np.divide(np.subtract(X*self.Shape, self.Coef), self.denomin)
        variable = self.__variable_instant_transform(X, self.Coef, self.denomin)

        return variable
    
    def __variable_instant_inverse_transform(self, rescaled_X):
        """
        inverse transformation of __variable_instant 
        """
        
        if self.with_variable_instant:
            unscaled_X = (rescaled_X*self.denomin+self.Coef)/self.Shape
        else: 
            unscaled_X = rescaled_X
        
        return unscaled_X
    
    def __compute_correlation_matrix(self, X, y):
        if hasattr(y, 'name'):
            name = y.name
        else :
            name=y.columns
        
        self.corr_X = pd.DataFrame(np.corrcoef(pd.concat([X, y], axis=1).T), columns= X.columns.tolist() + [name])
        return
        
    def __feature_selection(self, res, corr_X, M, res_list, threshold):
    
        #identifier la meilleure variable explicative
        corr_X= np.abs(corr_X)
        
        best_interaction = corr_X.iloc[-1, : -1].idxmax(axis=1)
        best_corr = M.loc[:, best_interaction]
        best_coef = corr_X.iloc[-1, : -1].max()

        if best_coef >= threshold : 
            #enregistrement du nom de la correlation et du coef dans la liste de reponses
            res_list.append([best_interaction, round(best_coef, 3), corr_X.columns.get_loc(best_interaction)])

            #Ajout de la colonne à la matrice des résultats
            res = pd.concat((res, best_corr), axis=1)

        return res, res_list
        
    
    def __partial_correlations(self, correlation_matrix, interaction):
    
        """
        Partial correlation Formula :
        rAB.C = rAB-rAC.rBC/(sqrt(1-rAC^2).sqrt(1-rBC^2)) 
        """

        reg = correlation_matrix.iloc[:, correlation_matrix.columns.get_loc(interaction)]

        rAC = np.array(reg).reshape(reg.shape[0], 1)
        rBC = np.array(reg).reshape(1, reg.shape[0])

        matrice = np.array(correlation_matrix)

        numerator = np.subtract(matrice, rAC.dot(rBC))

        #avoid negative values in sqrt
        rAC_sqr = np.square(rAC)
        rAC_sqr[rAC_sqr > 1]=0.99
        rBC_sqr = np.square(rBC)
        rBC_sqr[rBC_sqr > 1]=0.99

        denominator = np.sqrt(np.subtract(np.ones(rAC.shape),rAC_sqr)).dot(np.sqrt(np.subtract(np.ones(rBC.shape),rBC_sqr)))

        #avoid zero division
        denominator[denominator==0] = 1000


        correlation_matrix = pd.DataFrame(np.divide(numerator,denominator), columns=correlation_matrix.columns)

        for i,j in range(correlation_matrix.shape[0], correlation_matrix.shape[0]):
            correlation_matrix.iloc[i,j] = 1

        return correlation_matrix
    
    
    def __model_evaluation(self, mat_res):
    
        model = LinearRegression()

        #Calcul de Q2 global
        X = np.array(mat_res.iloc[:,1:])
        y = np.array(mat_res.iloc[:,0])
        
        loot=LeaveOneOut()
        loot.get_n_splits(X)

        SSres = []
        SStot = []

        #Leave-one-out cross validation
        for train_index, test_index in loot.split(X):
                X_train, X_test = X[train_index], X[test_index]
                y_train, y_test = y[train_index], y[test_index]
                model.fit(X_train, y_train)
                y_pred = model.predict(X_test)
                SSres.append((float(y_test[0])-float(y_pred[0]))**2)
                SStot.append(float(y_test[0]))


        SStot = np.array(SStot)
        SStot_mean = np.multiply(np.ones(SStot.shape), np.mean(SStot))    

        return round(1-np.sum(SSres)/np.sum(np.square(SStot-SStot_mean)),3)
    
    
    def __desirability_DS(self, target, target_weights, prediction):
        """
        Computation of desirability according to Derringer and Suich (1980) for multiple response optimization
        
        parameters:
            target : String, float, int or List of those
                    if string -> 'maximize' or 'minimize' or 'none'
                    float or int should correspond to the values that are targeted during optimization
                    List in the order of the columns in the dataFrame of responses
                    2 types of values are accepted : 
                                - String ('maximize' or 'minimize')
                                - Float (or int) if the optimization must target a specific value
        return
            DataFrame
            
        """
       
        desirability = pd.DataFrame(np.ones((prediction.shape[0], 1)), columns=['desirability']) 
                 
        for i in range(0, prediction.shape[1]):
            if prediction.shape[1] > 1:
                location = prediction.iloc[:, i]
            else:
                location = prediction
        
            if type(target) is str or type(target) is float or type(target) is int or type(target) is None:
                goal = target
            elif type(target) is list:
                goal= target[i]
            else:
                raise TypeError("target is either a string, a float, an integer or a list")

                    
            if goal=='maximize': #desirability to maximise the response
                objective = np.divide(location - location.min(axis = 0), location.max(axis=0)- location.min(axis=0))
            elif goal=='minimize': #desirability to minimize the response
                objective = np.divide(location.max(axis = 0) - location, location.max(axis=0) - location.min(axis=0))
            elif (goal == 'none') or (goal is None): #desirability to reach a specific target value
                objective = 0
            else:
                Solution1 = (location - location.min(axis=0))/ (goal - location.min(axis=0))
                Solution2 = (location - location.max(axis=0))/ (goal - location.max(axis=0))
                objective = np.minimum(Solution1, Solution2)
                    
            if target_weights is None:
                target_weights = [1 for n in range(0,prediction.shape[1])]
                    

            desirability = np.multiply(desirability, np.power(np.array(objective), target_weights[i]/np.sum(target_weights)).reshape(-1,1))
        return desirability
    
    def transform(self, X, y=None, scaler: str ='robust', variable_instant:bool=True, allow_autointeraction=False, 
                  interaction_list: list =['X_fort_Quand_Y_faible_Et_Inversement', 'X_fort_Ou_Y_fort', 'X_fort_Ou_Y_faible', 
                                    'X_et_Y_forts,X_fort_et_Y_faible', 'X_fort_si_Y_fort', 'X_fort_si_Y_faible', 
                                    'X_fort_si_Y_moyen', 'X_moyen_si_Y_fort', 'Ni_X_ni_Y_extremes', 'X_Y_moyens', 
                                   ' X_comme_Y', 'Somme_X_et_Y_forte', 'Difference_X_et_Y_forte']):
        """
        transform method :
        
        Params:
            X : DataFrame, matrix of the features
            y : DataFrame, matrix of the targer
            scaler : string, correspond to the method that will be used to rescale the data before computation of the interactions
            variable_instant : boolean, if True, data and computed interactions will be rescaled according to the "variable-instant" method of Lesty et al. (1999)
            allow_autointeraction : boolean, if True, additional interactions of the features with themselves will be considered
            interaction_list : list, List of interactions that are found relevant to the studied problem
        
        Return :
            self
        
        """
        if scaler not in ['robust', 'standard', 'minmax']:
            raise ValueError(f'{scaler} method is not implemented,\n  Implemented methods "robust, minmax and standard"')
        
        #retourne tableau de données avec interactions
        start = time.time()
        
        self.X = X.reset_index(drop=True)
        self.X_start = self.X
        self.y = y.reset_index(drop=True)
       
        #Step1: Rescale data
        try:
            self.X = pd.DataFrame(self.__rescale_data(self.X, self.y, scaler), columns = self.X.columns)
        except:
            raise NotImplementedError('rescaling data failed')
    
        #Step2: compute new features
        try:
            self.features = self.__compute_interaction(self.X, allow_autointeraction, interaction_list)
        except: 
            raise NotImplementedError('computation of the interactions failed')
        
        #Step3: Rescale data
        try:
            if variable_instant:
                try:
                    self.rescaled_features = pd.DataFrame(self.__variable_instant(self.features), columns = self.features.columns)
                    print('method = variable instant')
                except:
                    raise NotImplementedError('rescaling variable_instant data failed')
            else:
                try:
                    self.X = pd.DataFrame(self.__rescale_data(self.X, self.y, scaler), columns = self.X.columns)
                    print(f'method = {scaler}')
                except:
                    raise NotImplementedError(f'rescaling interactions with {scaler} method failed')
        except:
            raise NotImplementedError('rescaling data failed')
            
        end = time.time()
        print(f'calculé en {round(end-start, 3)} secondes')
        
        self.with_transform = True
        
        return self
    

    def fit(self, max_regressors_nb: int = 10, threshold: float = 0.2):
        
        """
        fit method :
        
        Params:
            max_regressors_nb : integer, set the maximum number of features that the model can use to describe the target
            threshold : float, threshold of partial correlation under which the model with stop including features into the model
        
        Return :
            self
        
        """
        
        #Mesure du temps de calcul
        start = time.time()
        
        #check if all parameters are correctly input
        if type(threshold) is not float:
            raise TypeError('threshold must be a float between 0 and 1')
        
        if type(max_regressors_nb) is not int:
            raise TypeError('max_regressors_nb must be an integer')
        
        self.model={}
        
        try:
            y = self.y.to_frame()
        except:
            y = self.y
        
        for i in y:
            self.model[i] = {}
            self.model[i]['selected_features'] = []
            self.model[i]['results'] = y[i]
            self.model[i]['metrics'] = []
            
            self.__compute_correlation_matrix(self.rescaled_features,  y[i])
            
            for reg in range(max_regressors_nb):
                #identification of the best interaction of this loop
                self.model[i]['results'], self.model[i]['selected_features'] = self.__feature_selection(self.model[i]['results'], self.corr_X, self.rescaled_features, self.model[i]['selected_features'], threshold)
                #compute the partial correlation
                self.corr_X = self.__partial_correlations(self.corr_X, self.model[i]['selected_features'][-1][0])
                #save values of Q² for model selection 
                self.model[i]['metrics'].append(self.__model_evaluation(self.model[i]['results']))

            #identify the best number of predictors by maximizing the Q² value
            self.model[i]['nb_predictor'] = self.model[i]['metrics'].index(max(self.model[i]['metrics']))+1
            #obsolete
            #self.model[i]['model_final'] = LinearRegression()

            #add a column of 1 (intercept) to the data for model fitting to responses using statsmodel.OLS    
            data = pd.concat((self.model[i]['results'].iloc[:,1:self.model[i]['nb_predictor']+1], pd.DataFrame(np.ones(y[i].shape), columns=['intercept'])), axis=1)
            #fitting predictors to responses with sm.OLS
            model = sm.OLS(y[i], data)
            self.model[i]['model_final'] = model.fit()
            #print results
            print(self.model[i]['model_final'].summary())
            
            #predict data
            self.model[i]['y_pred'] = self.model[i]['model_final'].predict(data)
            

        #print the amount of time spent to compute the model    
        end = time.time()
        print(f'fit method computed in {round(end-start, 3)} seconds')
        self.with_fit = True
        return self

    def fit_transform(self, X, y, **params):
        #select and pass kwargs to transform method
        transform_args = [key for key, value in inspect.signature(self.transform).parameters.items()]
        transform_dict = {key: params.pop(key) for key in dict(params) if key in transform_args}
        self.transform(X, y, **transform_dict)
        
        #select and pass kwargs to fit method
        fit_args = [key for key in inspect.signature(self.fit).parameters.items()]
        fit_dict = {key: params.pop(key) for key in dict(params) if key in fit_args}
        self.fit(**fit_dict)
        return self
    
    def predict(self, X): 
        #transformation of the matrix of parameters to predict
        X = X[self.X_start.columns]

        transformed_X = pd.DataFrame(self.transformer.transform(X), columns=X.columns.tolist())
        transformed_X_start = pd.DataFrame(self.transformer.transform(self.X_start), columns=X.columns.tolist())
        self.y_pred = pd.DataFrame()
        
        #transform y to DataFrame and avoid pandas.Series
        try:
            y = self.y.to_frame()
        except:
            y = self.y
        
        #loop through the models for each responses
        for i in y:
            new_X = None
        
            #transform the features to the selected predictors of the model in the fit method
            for element in self.model[i]['selected_features'][:self.model[i]['nb_predictor']]:
                try :
                    #call the right interaction
                    func = interaction_dict[element[0]]['interaction']
                    x_df= transformed_X[interaction_dict[element[0]]["x"].name]
                    y_df= transformed_X[interaction_dict[element[0]]["y"].name]

                    #call the max and min from the data on which the model was fitted
                    # otherwise if the new data to predict have not the same min and max for each feature the responses will be greatly modified
                    max_x = np.max(transformed_X_start[interaction_dict[element[0]]["x"].name])
                    min_x = np.min(transformed_X_start[interaction_dict[element[0]]["x"].name])
                    max_y = np.max(transformed_X_start[interaction_dict[element[0]]["y"].name])
                    min_y = np.min(transformed_X_start[interaction_dict[element[0]]["y"].name])

                    #compute the interaction values
                    col = eval(func)(x_df, y_df, max_x=max_x, min_x=min_x, max_y=max_y, min_y=min_y).compute()
                except KeyError:
                    #keep the feature unmodified if the modelisation found it relevant so
                    col = transformed_X[element[0]]
            
                finally:
                    #some 'infinity values' appear during calculation because of value very close to zero 
                    pd.options.mode.use_inf_as_na = True
                    
                    if any(col.isna()):
                        col[col.isna()] = 0
                    
                    #the new frame of predictors
                    if 'new_X' not in locals():
                        new_X = col
                    else:
                        #append the new calculated colums
                        new_X = pd.concat((new_X, col), axis=1)
            
            #transform the newly generated frame to acurate scale
            #L = mask to locate the position of the interaction in the Coef and denomin vectors
            self.L = [int(n) for n in np.array(self.model[i]['selected_features'][:self.model[i]['nb_predictor']])[:,-1]]

            #transform to the scale
            var_X =self.__variable_instant_transform(new_X, self.Coef, self.denomin, self.L)
            
            #transform to dataFrame
            variable_instant_X = pd.DataFrame(var_X, columns=self.model[i]['selected_features'][:self.model[i]['nb_predictor']])
            

            #compute with the model coefficients and intercept
            predictions = np.dot(variable_instant_X, self.model[i]['model_final'].params[:-1])
            predictions = predictions + self.model[i]['model_final'].params[-1]

            #save data in in y_pred varaible
            self.model[i]['y_pred'] = pd.DataFrame(predictions , columns=[f'Predicted {i}'])
            #Append the dataFrame with all the predicted response
            self.y_pred = pd.concat((self.y_pred, self.model[i]['y_pred']), axis=1)
        
        return self.y_pred
    
    def features_analysis(self, X):
        
        self.experimental_domain = {}
        
        for feature in X.columns.tolist():

            if len(X[feature].unique()) == len(set([int(num) for num in X[feature].values])):
                vartype = 'discrete'
                varlist = X[feature].unique().tolist()
            
            else:
                vartype = 'continuous'
                varlist = None
            
            #self.experimental_domain[X[feature].name] = [None, X[feature].min(axis=0), X[feature].max(axis=0), varlist , vartype]
            self.experimental_domain[X[feature].name] = {
                'min_value' : X[feature].min(axis=0),
                'max_value' : X[feature].max(axis=0),
                'var_type' : vartype,
                'list of values' : varlist,
                'fix_to_value' : None,
                'plot': False
                }

        
        self.mix=None
        for i in range(0, X.shape[1]):
            for j in range(i+1, X.shape[1]+1):
                a = X.iloc[:, i:j].sum(axis=1)
                
                if sum(np.abs(a - a.mean())) < sum(np.abs(a*0.05)):
                    self.mix = X.iloc[:, i:j].columns.tolist()
                    self.mixmax = X.iloc[:, i:j].max(axis=0).mean()
                    self.mixmin = X.iloc[:, i:j].min(axis=0).mean()
                    break
        
        #exp_dom = pd.DataFrame(self.experimental_domain, index=['status','min value', 'max_value', 'values', 'var type'])
        #print('experimental domain: ', exp_dom, 'mixture: ', self.mix, sep='\n\n' )

        return self.experimental_domain, self.mix

    def __features_generator(self, remaining_features, experimental_domain, size):
        #create random array respecting the restrictions of the features 
        for var in remaining_features:
            if experimental_domain[var]['var_type'] == 'discrete':
                #create a random array of the discrete values
                exploration_array = np.random.choice(experimental_domain[var]['list of values'], (size, 1), replace=True, p=None) #p peut permettre de mettre du poids sur le paramètre interessant
            elif experimental_domain[var]['var_type'] == 'continuous':
                #create a random array of the discrete values
                rng = np.random.default_rng()
                exploration_array = (experimental_domain[var]['max_value'] - experimental_domain[var]['min_value']) * rng.random((size, 1), dtype=np.float64) + experimental_domain[var]['min_value']
            experimental_domain[var]['generated values'] = exploration_array
        return experimental_domain
    
    def __mix_features_generator(self, alpha, size, random_state, mix):
        if alpha is None:
             alpha = np.ones((len(mix))) / len(mix)
        return pd.DataFrame(dirichlet.rvs(alpha, size=size, random_state=random_state), columns = mix)
    
    def generator(self, experimental_domain, mix, alpha : list, size: int, random_state: int=None):
        x=None
        
        if mix is not None:
                x = (self.mixmax- self.mixmin) * self.__mix_features_generator(alpha, size, random_state, mix) + self.mixmin
                
                remaining_features = list(set(self.X.columns.tolist()) - set(mix))
        else:
            remaining_features = experimental_domain.keys()
        print("size = " + str(size))   
        if len(remaining_features) > 0:
            experimental_domain = self.__features_generator(remaining_features, experimental_domain, size)
            for var in remaining_features:
                if x is None:
                    x = pd.DataFrame()   
                #a voir si erreur
                x = pd.concat((x, pd.DataFrame(experimental_domain[var]['generated values'], columns=[var])), axis=1)
        return x
    
    def optimize(self, experimental_domain:dict=None, target:list=None, target_weights:list=None, mix:list = None, alpha : list=None, size: int= 10000, random_state: int=None):
        
        #etude de la qualité des paramètres (quantitatif ou qualitatif)
        if experimental_domain is None:
            experimental_domain, mix = self.features_analysis(self.X_start)
              
        elif type(experimental_domain) is not dict:
            raise TypeError('experimental_domain must be a dictionary')
        
        #generation d'un nouveau tableau exploratoire
        sample = self.generator(experimental_domain, mix, alpha, size, random_state)

        #calcul des targets
        if self.with_transform:
            prediction = self.predict(sample)

            desirability = self.__desirability_DS(target, target_weights, prediction)

                #rendu
            res = pd.concat((sample, prediction, desirability), axis=1)
            
            output = pd.concat((np.around(res.nlargest(5, 'desirability').mean(axis=0), 3), np.around(res.iloc[res['desirability'].idxmax(axis=1), :], 3)), axis=1)
            output.columns = ['Mean of the 5 best results', 'Best result']
            
            print(output)
            
            self.with_optimization = True
            return res   
    
    def fitting_score(self, y):
        self.model[y.name]['r2score'] = r2_score(y, self.model[y.name]['y_pred']).round(3)
        self.model[y.name]['adjustedr2score'] = (1-(1-self.model[y.name]['r2score'])*(self.model[y.name]['results'].shape[0]-1)/(self.model[y.name]['results'].shape[0]-self.model[y.name]['nb_predictor']-1)).round(3)
        self.model[y.name]['Q2_obs'] = self.model[y.name]['metrics'][self.model[y.name]['nb_predictor']-1]
        stats = pd.DataFrame(np.array([self.model[y.name]['r2score'], self.model[y.name]['adjustedr2score'], self.model[y.name]['Q2_obs'].round(3)]).reshape(1,-1), columns=['R²', 'adj-R²','calc-Q²'], index = ['model score'])
        print(stats)
        return

    def residues_hist(self, y):
        plt.scatter(y, self.model[y.name]['y_pred']-y)
        plt.xlabel('residual distance')
        plt.ylabel('observation number')
        plt.title('Distribution of the residuals')
        return
    
    def fit_scatter(self, y):
        plt.scatter(y, self.model[y.name]['y_pred'])
        plt.plot([np.min(y), np.max(y)], [np.min(y),np.max(y)], c='r')
        plt.xlabel(f'measured values of {self.model[y.name]["results"].iloc[:,0].name}')
        plt.ylabel(f'predicted values of {self.model[y.name]["results"].iloc[:,0].name}')
        plt.title('Measured vs. predicted values')
        return
    
    def metrics_curve(self, y):
        plt.plot(np.arange(1, len(self.model[y.name]['metrics'])+1), self.model[y.name]['metrics'])
        plt.plot(self.model[y.name]['nb_predictor'], self.model[y.name]['metrics'][self.model[y.name]['nb_predictor']-1], linestyle='none', marker='o', label='Optimal number of predictors')
        plt.legend()
        plt.xlabel('number of predictors')
        plt.ylabel('Q² values')
        plt.title('Evolution of Q² with increasing number of predictors')
        return
    
    def describe(self, X, y):
        #retourne score, histogramme des résidus et diagonale
        try:
            y = self.y.to_frame()
        except:
            y = self.y
        
        for i in y:
            print(i)
            self.fitting_score(y[i])
            plt.figure()
            self.fit_scatter(y[i])
            plt.show()
            plt.figure()
            self.residues_hist(y[i])
            plt.show()
            plt.figure()
            self.metrics_curve(y[i])
            plt.show()
            #plt.suptitle('Overview of the modelisation')
        return

    
    def print_in_file(self, title: str ='title'):
                

        document = Document()

        document.add_heading(f'{title}', 0)
        if self.with_fit:
            document.add_heading('Modelization', level=1)
            document.add_heading('Model', level=2, style='List Number')
            document.add_heading('Metrics', level=2, style='List Number')
            document.add_heading('Plots', level=2, style='List Number')
            document.add_heading('Sensibility analysis', level=2, style='List Number')
    
            records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
            )

            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Qty'
            hdr_cells[1].text = 'Id'
            hdr_cells[2].text = 'Desc'
            for qty, id, desc in records:
                row_cells = table.add_row().cells
                row_cells[0].text = str(qty)
                row_cells[1].text = id
                row_cells[2].text = desc
            
            document.add_page_break()

        if self.with_optimization:
            document.add_heading('Optimization', level=1)
            document.add_heading('Targets', level=2, style='List Number')
            document.add_heading('Best Trials', level=2, style='List Number')
            document.add_heading('Pareto', level=2, style='List Number')
        

        document.add_picture('monty-truth.png', width=Inches(1.25))


        document.save(f'{title}.docx')

    def print_in_file(self):
        return #fichier avec données enregistrées et formatées

    def __extract_features(self, experimental_domain: dict):

        try:
            screened_var = [key for key in experimental_domain.keys() if experimental_domain[key]['plot'] == True]

            return screened_var

        except ValueError:
            print('To plot a ternary diagram please set 3 variable values to None')
        
    def __generate_ternary_matrix(self, experimental_domain, mix, alpha, size, random_state):
        #list the features to plot
        var = self.__extract_features(experimental_domain)
        #generate a dataset
        Arr = self.__mix_features_generator(alpha, size, random_state, var)
        
        #scale the data to the right values
        df_Arr = (self.mixmax * pd.DataFrame(Arr, columns=var) - self.mixmin)

        #list the feature not plot
        set_values = []
        set_values_names = []

        for key, value in experimental_domain.items():
            if isinstance(experimental_domain[key]['fix_to_value'], (int, float)):
                set_values.append(value['fix_to_value'])
                set_values_names.append(key)
            #if experimental_domain[key][0] is None:
                
                

        #broadcast the set values to complete the dataset and
        Bc_set_values = np.broadcast_to(np.array(set_values).reshape(1,-1), (size, len(set_values)))
        X = pd.DataFrame(np.hstack((df_Arr, Bc_set_values)), columns = var + set_values_names)

        """
        <!!> Probleme si toutes les variables ne font pas partie d'un plan de mélange !!!
        """
        #set the mixture to the acurate sum
        #print(X)
        X_mix = 100* X[mix] / np.sum(X[mix], axis=1).mean()
        X = pd.concat((X[set(X.columns)-set(X_mix.columns)], X_mix), axis=1)
        #print(X[self.X_start.columns])
        Results = self.predict(X[self.X_start.columns])

        return var, X, Results

    def TM_plot(self, experimental_domain: dict, mix: list = None, alpha: list=None, size: int = 1000, random_state: int=None, ncontours: int=20):
        #generate the accurate data to be plot
        if mix is None:
            if self.with_fit:
                mix = self.mix
        var, X, results = self.__generate_ternary_matrix(experimental_domain, mix, alpha, size, random_state)
        
        #plot the ternary contour for each targets
        for res in results:
            fig = ff.create_ternary_contour(np.array(X[var]).T, np.array(results[res]).T, pole_labels=var, interp_mode='cartesian', colorscale='Viridis', showscale=True, ncontours=ncontours)
            fig.show()
        return self
    
    def RS_plot(self, X: DataFrame = None, Y: DataFrame = None, experimental_domain: dict=None, status=None, size: int = 10000):
        #If variables are undefined by user, get the variables that were use for modelling
        if self.with_fit:
            if X is None:
                X = self.X_start 
            if Y is None:
                Y= self.y
            if experimental_domain is None:
                experimental_domain = self.experimental_domain
        else : 
            #if fit method has not been used yet -> user must define X and Y
            if (X is None) or (Y is None):
                raise ValueError('')
        
        #List the variables to plot in the model
        features_to_plot = self.__extract_features(experimental_domain)
        
        #generate the data that will be plotted
        X_complete = self.generator(experimental_domain= experimental_domain, mix= None, alpha= None, size=size)

        #Set the fixed variables to the desired value
        Arr=[]
        Arr_name=[]
        for key in (set(experimental_domain.keys())-set(features_to_plot)):
            if not isinstance(experimental_domain[key]['fix_to_value'], (int, float)):
                Arr.append(0)
            else:
                Arr.append(experimental_domain[key]['fix_to_value'])
            Arr_name.append(key)
        X_complete[Arr_name] = pd.DataFrame(np.full((size, len(Arr)), Arr), columns = Arr_name)
        
        #Compute the output values of the data
        Y = self.predict(X_complete[self.X_start.columns])

        #Plot the surface for each target
        a,b = X_complete[features_to_plot[0]], X_complete[features_to_plot[1]]
        for c in Y:

            fig = plt.figure(figsize=(14,9))    
            ax = plt.axes(projection='3d')
            Cmap = plt.get_cmap('viridis')

            Plot_surface(a.values.flatten(), b.values.flatten(), Y[c].values.flatten())
            
            """
            
            Z = Y[c].values.flatten()
            surf = ax.plot_trisurf(a.values.flatten(), b.values.flatten(), Z, cmap=Cmap, antialiased=True, edgecolor='none')
            fig.colorbar(surf, ax =ax, shrink=0.5, aspect=5)
            ax.set_xlabel(f'{a.name}')
            ax.set_ylabel(f'{b.name}')
            ax.set_zlabel(f'{c}')
            ax.set_title(f'{c}=f({a.name, b.name}')
                    
            fig.show()
            """
    
    def pareto_frontier(self, res, objectives: list, target: list = ['maximize', 'maximize'], plot: bool = True):

        """
        according to: Jamie Bull | jamiebull1@gmail.com
        https://oco-carbon.com/metrics/find-pareto-frontiers-in-python/
        """
        # Sort the list in either ascending or descending order of X
        if target[0]=='maximize':
            maxX = True
        elif target[0]=='minimize':
            maxX= False

        maxY = False if target[1] == 'minimize' else True
        
        if len(objectives) > 2:
            raise ValueError('Length of "objectives" should be 2.')

        myList = sorted([[res[objectives[0]][i], res[objectives[1]][i]] for i in range(len(res[objectives[0]]))], reverse=maxX)
        # Start the Pareto frontier with the first value in the sorted list
        p_front = [myList[0]]    
        # Loop through the sorted list
        for pair in myList[1:]:
            if maxY: 
                if pair[1] >= p_front[-1][1]: # Look for higher values of Y…
                    p_front.append(pair) # … and add them to the Pareto frontier
            else:
                if pair[1] <= p_front[-1][1]: # Look for lower values of Y…
                    p_front.append(pair) # … and add them to the Pareto frontier
        # Turn resulting pairs back into a list of Xs and Ys
        p_frontX = [pair[0] for pair in p_front]
        p_frontY = [pair[1] for pair in p_front]

        mask = res[objectives[0]].isin(p_frontX)

        if plot:
            plt.scatter(res[objectives[0]], res[objectives[1]], alpha=0.5, c='lightgrey', label='trials')
            plt.scatter(res[objectives[0]].mask(~mask), res[objectives[1]].mask(~mask), label='undominated trials')
            # Then plot the Pareto frontier on top
            plt.plot(p_frontX, p_frontY, c='r', label='pareto front')
            plt.xlabel(res[objectives[0]].name)
            plt.ylabel(res[objectives[1]].name)
            plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', borderaxespad=0.)
            plt.show()

        return p_frontX, p_frontY
    
    def sensibility_analysis(self, experimental_domain: dict, plot: bool = True):
        
        Sobol_list = []

        problem = {
            'num_vars': len(self.X.columns),
            'names' : list(self.X.columns),
            'bounds': [[experimental_domain[n]['min_value'], experimental_domain[n]['max_value']] for n in self.X.columns]
        }
        print(problem)
        param_values = pd.DataFrame(saltelli.sample(problem, 1024), columns=self.X.columns)
        print(param_values)
        predictions = self.predict(param_values)
        
        for c in predictions:
            Si = sobol.analyze(problem, np.array(predictions[c]))
            Sobol_list.append(Si)
            t, fo, so = Si.to_df()

            fo_name = [str(n) for n in fo.index.tolist()]
            so_name = [str(n) for n in so.index.tolist()]
            t_name = [str(n) for n in t.index.tolist()]

            if plot:
                fig, (ax1, ax2, ax3) = plt.subplots(3, sharex =True)
                ax1.barh(fo_name,fo['S1'].values.round(3) , 0.5, color='cornflowerblue', label='Parameters', xerr=fo['S1_conf'].values)
                ax2.barh(so_name, so['S2'].values.round(3), 0.5, color='lightsteelblue', label='Interactions', xerr=so['S2_conf'].values)
                ax3.barh(t_name,t['ST'].values.round(3), 0.5, color='royalblue', label='Total', xerr=(t['ST_conf'].values))
                fig.legend(bbox_to_anchor=(1.05, 0.85), loc='upper left', borderaxespad=0.)
                plt.show()

        return Sobol_list

    
    def outliers_influence(self, plot: bool =True):
        frame_list=[]
        for i in self.y:
            outliers = outliers_influence.OLSInfluence(self.model[i]['model_final'])
            frame_list.append(outliers.summary_frame())
            threshold = 4/outliers.summary_frame().shape[0]
            
            print(f'threshold (4/n) = {round(threshold,3)}' )
            outliers_list= []
            for n in range(0,outliers.summary_frame().shape[0]):
                if outliers.summary_frame().at[n, 'cooks_d'] >= threshold:
                    outliers_list.append((n, outliers.summary_frame().at[n, 'cooks_d']))
            print(f'potential outliers : {outliers_list}')
            if plot:
                #plt.figure()
                plt.scatter(range(0,outliers.summary_frame().shape[0]), outliers.summary_frame()['cooks_d'], label=i)
            print(outliers.summary_table())
            
        plt.plot([0, outliers.summary_frame().shape[0]], [threshold, threshold], c='r', label='threshold')
        plt.xlabel('Observation indices')
        plt.ylabel('Cook\'s distance')
        plt.legend(bbox_to_anchor=(1.05, 0.85), loc='upper left', borderaxespad=0.)
        plt.show()
        
        return frame_list
